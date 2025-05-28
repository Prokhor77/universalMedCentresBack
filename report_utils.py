import os
import smtplib
import urllib
from email import encoders
from email.header import Header
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
import docx
import transliterate
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import sqlite3
from fpdf import FPDF

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)


def generate_report_docx(med_center_id, period_days, db_connection):

    end_date = datetime.now()
    start_date = end_date - timedelta(days=period_days)

    # Format dates for display and queries
    end_date_str = end_date.strftime("%Y-%m-%d")
    start_date_str = start_date.strftime("%Y-%m-%d")

    # Create a new document
    doc = docx.Document()

    # Add title
    title = doc.add_heading('Медицинский отчет', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add report period
    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_para.add_run(f'Период: {start_date.strftime("%d.%m.%Y")} - {end_date.strftime("%d.%m.%Y")}')
    period_run.font.size = Pt(12)

    # Get medical center info
    cursor = db_connection.cursor()
    cursor.execute("SELECT centerName, centerAddress, centerNumber FROM med_centers WHERE idCenter = ?",
                   (med_center_id,))
    center = cursor.fetchone()

    if center:
        # Add medical center info
        doc.add_heading('Информация о медицинском центре', level=1)
        doc.add_paragraph(f'Название: {center[0]}')
        doc.add_paragraph(f'Адрес: {center[1]}')
        doc.add_paragraph(f'Телефон: {center[2]}')

    # 1. General statistics
    doc.add_heading('Общая статистика', level=1)

    # 1.1 Appointments count
    cursor.execute("""
        SELECT COUNT(*) FROM records 
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
    """, (med_center_id, start_date_str, end_date_str))
    appointments_count = cursor.fetchone()[0]
    doc.add_paragraph(f'Всего приемов за период: {appointments_count}')

    # 1.2 Doctors count
    cursor.execute("""
        SELECT COUNT(*) FROM users 
        WHERE medCenterId = ? AND role = 'doctor'
    """, (med_center_id,))
    doctors_count = cursor.fetchone()[0]
    doc.add_paragraph(f'Количество врачей: {doctors_count}')

    # 1.3 Inpatient count
    cursor.execute("""
        SELECT COUNT(*) FROM inpatient_care 
        WHERE medCenterId = ? AND active = 'true'
    """, (med_center_id,))
    inpatient_count = cursor.fetchone()[0]
    doc.add_paragraph(f'Количество пациентов стационарного лечения: {inpatient_count}')

    # 2. Financial statistics
    doc.add_heading('Финансовая статистика', level=1)

    # 2.1 Total income
    cursor.execute("""
        SELECT SUM(price) FROM records 
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
    """, (med_center_id, start_date_str, end_date_str))
    total_income = cursor.fetchone()[0] or 0
    doc.add_paragraph(f'Общий доход за период: {total_income} BYN')

    # 2.2 Paid vs Free services
    cursor.execute("""
        SELECT paidOrFree, COUNT(*) FROM records 
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
        GROUP BY paidOrFree
    """, (med_center_id, start_date_str, end_date_str))
    services = cursor.fetchall()

    paid_count = 0
    free_count = 0
    for service in services:
        if service[0] == 'payed':
            paid_count = service[1]
        elif service[0] == 'free':
            free_count = service[1]

    doc.add_paragraph(f'Платные услуги: {paid_count}')
    doc.add_paragraph(f'Бесплатные услуги: {free_count}')

    # 3. Doctor statistics
    doc.add_heading('Статистика по врачам', level=1)

    # 3.1 Top doctors by appointments
    cursor.execute("""
        SELECT u.fullName, d.work_type, COUNT(r.id) as appointment_count
        FROM records r
        JOIN users u ON r.doctorId = u.id
        JOIN doctors d ON u.id = d.userId
        WHERE r.medCenterId = ? AND r.time_end BETWEEN ? AND ?
        GROUP BY r.doctorId
        ORDER BY appointment_count DESC
        LIMIT 5
    """, (med_center_id, start_date_str, end_date_str))
    top_doctors = cursor.fetchall()

    if top_doctors:
        doc.add_paragraph('Топ врачей по количеству приемов:')
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'ФИО врача'
        header_cells[1].text = 'Специализация'
        header_cells[2].text = 'Количество приемов'

        # Add data rows
        for doctor in top_doctors:
            row_cells = table.add_row().cells
            row_cells[0].text = doctor[0]
            row_cells[1].text = doctor[1] or 'Не указана'
            row_cells[2].text = str(doctor[2])

    # 3.2 Average doctor rating
    cursor.execute("""
        SELECT u.fullName, d.work_type, AVG(f.grade) as avg_rating
        FROM feedback f
        JOIN doctors d ON f.doctorId = d.userId
        JOIN users u ON d.userId = u.id
        WHERE u.medCenterId = ? AND f.active = 'true'
        GROUP BY f.doctorId
        ORDER BY avg_rating DESC
    """, (med_center_id,))
    doctor_ratings = cursor.fetchall()

    if doctor_ratings:
        doc.add_paragraph('')  # Add some space
        doc.add_paragraph('Рейтинг врачей:')
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'ФИО врача'
        header_cells[1].text = 'Специализация'
        header_cells[2].text = 'Средний рейтинг'

        # Add data rows
        for doctor in doctor_ratings:
            row_cells = table.add_row().cells
            row_cells[0].text = doctor[0]
            row_cells[1].text = doctor[1] or 'Не указана'
            row_cells[2].text = f"{doctor[2]:.1f}" if doctor[2] else 'Нет отзывов'

    # 4. Feedback statistics
    doc.add_heading('Статистика отзывов', level=1)

    # 4.1 Feedback count by status
    cursor.execute("""
        SELECT active, COUNT(*) FROM feedback
        WHERE medCenterId = ?
        GROUP BY active
    """, (med_center_id,))
    feedback_stats = cursor.fetchall()

    approved_count = 0
    rejected_count = 0
    pending_count = 0

    for stat in feedback_stats:
        if stat[0] == 'true':
            approved_count = stat[1]
        elif stat[0] == 'false':
            rejected_count = stat[1]
        elif stat[0] == 'in_progress':
            pending_count = stat[1]

    doc.add_paragraph(f'Одобренные отзывы: {approved_count}')
    doc.add_paragraph(f'Отклоненные отзывы: {rejected_count}')
    doc.add_paragraph(f'Ожидающие модерации: {pending_count}')

    # 5. Daily statistics for the period
    doc.add_heading('Ежедневная статистика за период', level=1)

    # Create a table for daily stats
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Дата'
    header_cells[1].text = 'Количество приемов'
    header_cells[2].text = 'Доход (BYN)'

    # Get daily stats
    current_date = start_date
    while current_date <= end_date:
        date_str = current_date.strftime("%Y-%m-%d")

        # Get appointments count
        cursor.execute("""
            SELECT COUNT(*), SUM(price) FROM records 
            WHERE medCenterId = ? AND time_end LIKE ?
        """, (med_center_id, f"{date_str}%"))
        result = cursor.fetchone()
        daily_count = result[0] or 0
        daily_income = result[1] or 0

        # Add row to table
        row_cells = table.add_row().cells
        row_cells[0].text = current_date.strftime("%d.%m.%Y")
        row_cells[1].text = str(daily_count)
        row_cells[2].text = str(daily_income)

        current_date += timedelta(days=1)

    doc.add_heading('Детализация всех приемов за период', level=1)

    cursor.execute("""
        SELECT
            r.time_start,
            r.time_end,
            u.fullName AS patient_name,
            d_user.fullName AS doctor_name,
            d.work_type AS doctor_specialization,
            r.description,
            r.assignment,
            r.paidOrFree,
            r.price
        FROM records r
        LEFT JOIN users u ON r.userId = u.id
        LEFT JOIN users d_user ON r.doctorId = d_user.id
        LEFT JOIN doctors d ON d.userId = r.doctorId
        WHERE r.medCenterId = ? AND r.time_end BETWEEN ? AND ?
        ORDER BY r.time_end DESC
    """, (med_center_id, start_date_str, end_date_str))
    all_appointments = cursor.fetchall()

    if all_appointments:
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Дата/Время начала'
        hdr_cells[1].text = 'Дата/Время конца'
        hdr_cells[2].text = 'Пациент'
        hdr_cells[3].text = 'Врач'
        hdr_cells[4].text = 'Специализация'
        hdr_cells[5].text = 'Описание'
        hdr_cells[6].text = 'Назначения'
        hdr_cells[7].text = 'Тип услуги'
        hdr_cells[8].text = 'Стоимость'

        for row in all_appointments:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0] or ''
            row_cells[1].text = row[1] or ''
            row_cells[2].text = row[2] or ''
            row_cells[3].text = row[3] or ''
            row_cells[4].text = row[4] or ''
            row_cells[5].text = row[5] or ''
            row_cells[6].text = row[6] or ''
            row_cells[7].text = 'Платно' if row[7] == 'payed' else 'Бесплатно'
            row_cells[8].text = str(row[8]) if row[8] is not None else ''
    else:
        doc.add_paragraph('Нет приемов за выбранный период.')

    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{REPORTS_DIR}/report_{med_center_id}_{period_days}days_{timestamp}.docx"
    doc.save(filename)

    return filename


def generate_report_pdf(med_center_id, period_days, db_connection):
    """Generate a comprehensive report in PDF format"""
    # Calculate date range
    end_date = datetime.now()
    start_date = end_date - timedelta(days=period_days)

    # Format dates for display and queries
    end_date_str = end_date.strftime("%Y-%m-%d")
    start_date_str = start_date.strftime("%Y-%m-%d")

    # Create PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 14)

    # Add title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "Medical Report", ln=True, align="C")

    # Add report period
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, f"Period: {start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}", ln=True, align="C")

    # Get medical center info
    cursor = db_connection.cursor()
    cursor.execute("SELECT centerName, centerAddress, centerNumber FROM med_centers WHERE idCenter = ?",
                   (med_center_id,))
    center = cursor.fetchone()

    if center:
        # Add medical center info
        pdf.set_font("Arial", "B", 14)
        pdf.ln(5)
        pdf.cell(0, 10, "Medical Center Information", ln=True)

        pdf.set_font("Arial", "", 12)
        # Clean and transliterate data
        center_name_latin = clean_for_pdf(transliterate.translit(center[0], 'ru', reversed=True))
        center_address_latin = clean_for_pdf(transliterate.translit(center[1], 'ru', reversed=True))
        center_phone = clean_for_pdf(center[2])

        pdf.cell(0, 8, f"Name: {center_name_latin}", ln=True)
        pdf.cell(0, 8, f"Address: {center_address_latin}", ln=True)
        pdf.cell(0, 8, f"Phone: {center_phone}", ln=True)

    # 1. General statistics
    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "General Statistics", ln=True)

    pdf.set_font("Arial", "", 12)

    # 1.1 Appointments count
    cursor.execute("""
        SELECT COUNT(*) FROM records
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
    """, (med_center_id, start_date_str, end_date_str))
    appointments_count = cursor.fetchone()[0]
    pdf.cell(0, 8, f"Total appointments: {appointments_count}", ln=True)

    # 1.2 Doctors count
    cursor.execute("""
        SELECT COUNT(*) FROM users
        WHERE medCenterId = ? AND role = 'doctor'
    """, (med_center_id,))
    doctors_count = cursor.fetchone()[0]
    pdf.cell(0, 8, f"Number of doctors: {doctors_count}", ln=True)

    # 1.3 Inpatient count
    cursor.execute("""
        SELECT COUNT(*) FROM inpatient_care
        WHERE medCenterId = ? AND active = 'true'
    """, (med_center_id,))
    inpatient_count = cursor.fetchone()[0]
    pdf.cell(0, 8, f"Number of inpatients: {inpatient_count}", ln=True)

    # 2. Financial statistics
    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "Financial Statistics", ln=True)

    pdf.set_font("Arial", "", 12)

    # 2.1 Total income
    cursor.execute("""
        SELECT SUM(price) FROM records
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
    """, (med_center_id, start_date_str, end_date_str))
    total_income = cursor.fetchone()[0] or 0
    pdf.cell(0, 8, f"Total income for period: {total_income} BYN", ln=True)

    # 2.2 Paid vs Free services
    cursor.execute("""
        SELECT paidOrFree, COUNT(*) FROM records
        WHERE medCenterId = ? AND time_end BETWEEN ? AND ?
        GROUP BY paidOrFree
    """, (med_center_id, start_date_str, end_date_str))
    services = cursor.fetchall()

    paid_count = 0
    free_count = 0
    for service in services:
        if service[0] == 'payed':
            paid_count = service[1]
        elif service[0] == 'free':
            free_count = service[1]

    pdf.cell(0, 8, f"Paid services: {paid_count}", ln=True)
    pdf.cell(0, 8, f"Free services: {free_count}", ln=True)

    # 3. Doctor statistics
    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "Doctor Statistics", ln=True)

    pdf.set_font("Arial", "", 12)

    # 3.1 Top doctors by appointments
    cursor.execute("""
        SELECT u.fullName, d.work_type, COUNT(r.id) as appointment_count
        FROM records r
        JOIN users u ON r.doctorId = u.id
        JOIN doctors d ON u.id = d.userId
        WHERE r.medCenterId = ? AND r.time_end BETWEEN ? AND ?
        GROUP BY r.doctorId
        ORDER BY appointment_count DESC
        LIMIT 5
    """, (med_center_id, start_date_str, end_date_str))
    top_doctors = cursor.fetchall()

    if top_doctors:
        pdf.cell(0, 8, "Top doctors by number of appointments:", ln=True)

        # Table header
        pdf.set_font("Arial", "B", 12)
        pdf.cell(70, 8, "Doctor Name", border=1)
        pdf.cell(70, 8, "Specialization", border=1)
        pdf.cell(50, 8, "Appointments", border=1, ln=True)

        # Table data
        pdf.set_font("Arial", "", 12)
        for doctor in top_doctors:
            doctor_name = doctor[0] or ""
            specialization = doctor[1] or "Not specified"

            # Clean and transliterate data
            doctor_name_latin = clean_for_pdf(transliterate.translit(doctor_name, 'ru', reversed=True))
            specialization_latin = clean_for_pdf(transliterate.translit(specialization, 'ru', reversed=True))

            pdf.cell(70, 8, doctor_name_latin, border=1)
            pdf.cell(70, 8, specialization_latin, border=1)
            pdf.cell(50, 8, str(doctor[2]), border=1, ln=True)

    # 3.2 Average doctor rating
    cursor.execute("""
        SELECT u.fullName, d.work_type, AVG(f.grade) as avg_rating
        FROM feedback f
        JOIN doctors d ON f.doctorId = d.userId
        JOIN users u ON d.userId = u.id
        WHERE u.medCenterId = ? AND f.active = 'true'
        GROUP BY f.doctorId
        ORDER BY avg_rating DESC
    """, (med_center_id,))
    doctor_ratings = cursor.fetchall()

    if doctor_ratings:
        pdf.ln(5)
        pdf.cell(0, 8, "Doctor Ratings:", ln=True)

        # Table header
        pdf.set_font("Arial", "B", 12)
        pdf.cell(70, 8, "Doctor Name", border=1)
        pdf.cell(70, 8, "Specialization", border=1)
        pdf.cell(50, 8, "Average Rating", border=1, ln=True)

        # Table data
        pdf.set_font("Arial", "", 12)
        for doctor in doctor_ratings:
            doctor_name = doctor[0] or ""
            specialization = doctor[1] or "Not specified"

            # Clean and transliterate data
            doctor_name_latin = clean_for_pdf(transliterate.translit(doctor_name, 'ru', reversed=True))
            specialization_latin = clean_for_pdf(transliterate.translit(specialization, 'ru', reversed=True))

            rating_text = f"{doctor[2]:.1f}" if doctor[2] else "No reviews"
            pdf.cell(70, 8, doctor_name_latin, border=1)
            pdf.cell(70, 8, specialization_latin, border=1)
            pdf.cell(50, 8, rating_text, border=1, ln=True)

    # 4. Feedback statistics
    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "Feedback Statistics", ln=True)

    pdf.set_font("Arial", "", 12)

    # 4.1 Feedback count by status
    cursor.execute("""
        SELECT active, COUNT(*) FROM feedback
        WHERE medCenterId = ?
        GROUP BY active
    """, (med_center_id,))
    feedback_stats = cursor.fetchall()

    approved_count = 0
    rejected_count = 0
    pending_count = 0

    for stat in feedback_stats:
        if stat[0] == 'true':
            approved_count = stat[1]
        elif stat[0] == 'false':
            rejected_count = stat[1]
        elif stat[0] == 'in_progress':
            pending_count = stat[1]

    pdf.cell(0, 8, f"Approved reviews: {approved_count}", ln=True)
    pdf.cell(0, 8, f"Rejected reviews: {rejected_count}", ln=True)
    pdf.cell(0, 8, f"Pending moderation: {pending_count}", ln=True)

    # 5. Daily statistics for the period
    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "Daily Statistics for Period", ln=True)

    # Table header
    pdf.set_font("Arial", "B", 12)
    pdf.cell(60, 8, "Date", border=1)
    pdf.cell(60, 8, "Appointments", border=1)
    pdf.cell(60, 8, "Income (BYN)", border=1, ln=True)

    # Get daily stats
    pdf.set_font("Arial", "", 12)
    current_date = start_date
    while current_date <= end_date:
        date_str = current_date.strftime("%Y-%m-%d")

        # Get appointments count
        cursor.execute("""
            SELECT COUNT(*), SUM(price) FROM records
            WHERE medCenterId = ? AND time_end LIKE ?
        """, (med_center_id, f"{date_str}%"))
        result = cursor.fetchone()
        daily_count = result[0] or 0
        daily_income = result[1] or 0

        # Add row to table
        pdf.cell(60, 8, current_date.strftime("%d.%m.%Y"), border=1)
        pdf.cell(60, 8, str(daily_count), border=1)
        pdf.cell(60, 8, str(daily_income), border=1, ln=True)

        current_date += timedelta(days=1)

    cursor.execute("""
        SELECT
            r.time_start,
            r.time_end,
            u.fullName AS patient_name,
            d_user.fullName AS doctor_name,
            d.work_type AS doctor_specialization,
            r.description,
            r.assignment,
            r.paidOrFree,
            r.price
        FROM records r
        LEFT JOIN users u ON r.userId = u.id
        LEFT JOIN users d_user ON r.doctorId = d_user.id
        LEFT JOIN doctors d ON d.userId = r.doctorId
        WHERE r.medCenterId = ? AND r.time_end BETWEEN ? AND ?
        ORDER BY r.time_end DESC
    """, (med_center_id, start_date_str, end_date_str))
    all_appointments = cursor.fetchall()

    pdf.set_font("Arial", "B", 14)
    pdf.ln(5)
    pdf.cell(0, 10, "All Appointments for the Period", ln=True)

    if all_appointments:
        # Таблица: заголовки
        pdf.set_font("Arial", "B", 8)
        headers = [
            "Start", "End", "Patient", "Doctor", "Specialization",
            "Description", "Assignment", "Type", "Price"
        ]
        col_widths = [22, 22, 25, 25, 22, 32, 32, 13, 13]  # Можно подогнать под свой вкус

        for i, header in enumerate(headers):
            pdf.cell(col_widths[i], 7, header, border=1)
        pdf.ln()

        def safe_pdf(text):
            if not text:
                return ""
            try:
                return clean_for_pdf(transliterate.translit(str(text), 'ru', reversed=True))
            except Exception:
                return clean_for_pdf(str(text))

        # Данные
        pdf.set_font("Arial", "", 7)
        for row in all_appointments:
            def short(text, maxlen):
                t = safe_pdf(text)
                return (t[:maxlen - 3] + "...") if t and len(t) > maxlen else (t or "")

            pdf.cell(col_widths[0], 6, short(row[0], 16), border=1)
            pdf.cell(col_widths[1], 6, short(row[1], 16), border=1)
            pdf.cell(col_widths[2], 6, short(row[2], 15), border=1)
            pdf.cell(col_widths[3], 6, short(row[3], 15), border=1)
            pdf.cell(col_widths[4], 6, short(row[4], 12), border=1)
            pdf.cell(col_widths[5], 6, short(row[5], 20), border=1)
            pdf.cell(col_widths[6], 6, short(row[6], 20), border=1)
            pdf.cell(col_widths[7], 6, "Paid" if row[7] == "payed" else "Free", border=1)
            pdf.cell(col_widths[8], 6, str(row[8]) if row[8] is not None else "", border=1)
            pdf.ln()
    else:
        pdf.set_font("Arial", "", 12)
        pdf.cell(0, 8, "No appointments for the selected period.", ln=True)

    # Save the PDF
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = clean_for_pdf(transliterate.translit(center[0], 'ru', reversed=True)) if center else str(med_center_id)
    filename = f"{REPORTS_DIR}/report_{safe_name}_{period_days}days_{timestamp}.pdf"
    pdf.output(filename)

    return filename

def clean_for_pdf(text):
    if not text:
        return ""
    # Заменяем специальные символы на их ASCII эквиваленты
    replacements = {
        '№': 'No.',
        '©': '(c)',
        '®': '(R)',
        '™': '(TM)',
        '–': '-',
        '—': '-',
        '«': '"',
        '»': '"',
        '…': '...',
        '•': '*',
        '₽': 'RUB',
        '€': 'EUR',
        '£': 'GBP',
        '°': ' degrees',
        '±': '+/-',
        '×': 'x',
        '÷': '/',
        '≤': '<=',
        '≥': '>=',
        '≠': '!=',
        '≈': '~=',
        '∞': 'infinity'
    }
    for char, replacement in replacements.items():
        if char in text:
            text = text.replace(char, replacement)
    return text

def send_report_email(to_email, report_file_path, period_days):
    """Send the generated report via email"""
    username = "prohor.odinets@yandex.by"
    password = "hgrgaosbzvtxdxam"
    smtp_server = "smtp.yandex.com"
    smtp_port = 587

    # English period text
    period_text = ""
    if period_days == 1:
        period_text = "for the last day"
    elif period_days == 7:
        period_text = "for the last week"
    elif period_days == 30:
        period_text = "for the last month"

    msg = MIMEMultipart()
    msg['From'] = username
    msg['To'] = to_email
    msg['Subject'] = Header(f"Medical Report {period_text}", 'utf-8')

    html = f"""
    <html>
    <body>
        <h2>Medical Report {period_text}</h2>
        <p>Dear User,</p>
        <p>Attached is the requested report on the activities of the medical center {period_text}.</p>
        <p>Best regards,<br>Medical Center Management System</p>
    </body>
    </html>
    """
    msg.attach(MIMEText(html, 'html'))

    # Attach the report file
    with open(report_file_path, "rb") as f:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        filename = os.path.basename(report_file_path)
        try:
            filename.encode('ascii')
            part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        except UnicodeEncodeError:
            filename_utf8 = urllib.parse.quote(filename)
            part.add_header('Content-Disposition', f"attachment; filename*=UTF-8''{filename_utf8}")
        msg.attach(part)

    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(username, password)
        server.sendmail(username, to_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        print(f"Email sending error: {e}")
        return False


print("report_utils.py created successfully")